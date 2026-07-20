from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path

import pytesseract
from docx import Document
from pdf2image import convert_from_bytes
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}

MAX_FILE_SIZE = 8 * 1024 * 1024
MIN_EXTRACTED_CHARACTERS = 100
MAX_PDF_PAGES = 10
# Lower DPI keeps OCR image rendering within memory on a CPU box.
OCR_DPI = 150

_tesseract_checked = False
_tesseract_ok = False


def tesseract_available() -> bool:
    """Whether the Tesseract OCR binary is installed (checked once, cached)."""
    global _tesseract_checked, _tesseract_ok
    if not _tesseract_checked:
        try:
            pytesseract.get_tesseract_version()
            _tesseract_ok = True
        except Exception:
            _tesseract_ok = False
        _tesseract_checked = True
    return _tesseract_ok


@dataclass
class ExtractionResult:
    text: str
    method: str
    file_type: str


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    # Remove excessive spaces while preserving line breaks.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def preprocess_image(image: Image.Image) -> Image.Image:
    image = ImageOps.exif_transpose(image)
    image = image.convert("L")

    # Increase size for small CV screenshots.
    if image.width < 1600:
        scale = 1600 / image.width
        new_size = (
            int(image.width * scale),
            int(image.height * scale),
        )
        image = image.resize(new_size)

    image = ImageEnhance.Contrast(image).enhance(1.7)
    image = image.filter(ImageFilter.SHARPEN)

    return image


def run_ocr(image: Image.Image) -> str:
    processed = preprocess_image(image)

    try:
        return pytesseract.image_to_string(
            processed,
            lang="eng",
            config="--oem 3 --psm 6",
        )
    except Exception as error:
        raise ValueError(
            "OCR is not available on the server. Please paste the CV text "
            "or upload a text-based PDF or DOCX instead."
        ) from error


def extract_pdf_text(file_bytes: bytes) -> ExtractionResult:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as error:
        raise ValueError(
            "The PDF could not be read. It may be corrupted or password-protected."
        ) from error

    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError(
            f"PDF exceeds the maximum of {MAX_PDF_PAGES} pages."
        )

    page_text: list[str] = []

    for page in reader.pages:
        page_text.append(page.extract_text() or "")

    extracted = clean_text("\n\n".join(page_text))

    if len(extracted) >= MIN_EXTRACTED_CHARACTERS:
        return ExtractionResult(
            text=extracted,
            method="pypdf",
            file_type="pdf",
        )

    # The PDF is likely scanned. Fail fast if OCR is unavailable rather than
    # rendering page images (memory-heavy) and then crashing.
    if not tesseract_available():
        raise ValueError(
            "This PDF has little selectable text and OCR is not available on "
            "the server. Please paste the CV text or upload a text-based PDF "
            "or DOCX instead."
        )

    images = convert_from_bytes(
        file_bytes,
        dpi=OCR_DPI,
        fmt="png",
        first_page=1,
        last_page=MAX_PDF_PAGES,
    )

    try:
        ocr_pages = [run_ocr(image) for image in images]
    finally:
        for image in images:
            image.close()

    ocr_text = clean_text("\n\n".join(ocr_pages))

    return ExtractionResult(
        text=ocr_text,
        method="pdf-ocr",
        file_type="pdf",
    )


def extract_docx_text(file_bytes: bytes) -> ExtractionResult:
    document = Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # Some CV templates place content inside tables.
    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if cells:
                parts.append(" | ".join(cells))

    return ExtractionResult(
        text=clean_text("\n".join(parts)),
        method="python-docx",
        file_type="docx",
    )


def extract_image_text(
    file_bytes: bytes,
    extension: str,
) -> ExtractionResult:
    if not tesseract_available():
        raise ValueError(
            "OCR is not available on the server, so image CVs cannot be read. "
            "Please paste the CV text instead."
        )

    try:
        image = Image.open(io.BytesIO(file_bytes))
    except Exception as error:
        raise ValueError(
            "The image could not be read. Please upload a valid PNG, JPG, or JPEG."
        ) from error

    extracted = clean_text(run_ocr(image))

    return ExtractionResult(
        text=extracted,
        method="image-ocr",
        file_type=extension.lstrip("."),
    )


def extract_document_text(
    filename: str,
    file_bytes: bytes,
) -> ExtractionResult:
    if not filename:
        raise ValueError("The uploaded file has no filename.")

    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Unsupported file. Use PDF, DOCX, PNG, JPG, JPEG, or WEBP."
        )

    if not file_bytes:
        raise ValueError("The uploaded file is empty.")

    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError(
            "The file is too large. Maximum size is 8 MB."
        )

    if extension == ".pdf":
        result = extract_pdf_text(file_bytes)

    elif extension == ".docx":
        result = extract_docx_text(file_bytes)

    else:
        result = extract_image_text(
            file_bytes,
            extension,
        )

    if len(result.text) < MIN_EXTRACTED_CHARACTERS:
        raise ValueError(
            "Not enough readable CV text was found. "
            "Upload a clearer document or paste the CV text."
        )

    return result